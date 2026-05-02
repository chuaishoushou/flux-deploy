#!/usr/bin/env python3
"""
flux-deploy 使用手册 docx 后处理。

由 build-manual.sh 在 convert_style.py 之后调用。做三件事：

1. 横向 -> 纵向：FLUX 标准模板默认 A4 横向，对手册类文档太宽，翻成纵向。
2. F0/F1/F2 样式 jc=both -> jc=left：模板默认两端对齐，长路径会被强行拉伸成稀疏字距。
3. 识别"代码/命令/路径"段落 -> Consolas 等宽字体 + 浅灰底：还原 generate-manual.js 中
   被 convert_style.py 抹掉的代码块视觉。

幂等：重复跑无副作用。

用法: python3 postprocess-manual.py <docx 路径>
"""
import sys
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


CMD_STARTS = (
    "java -jar", "./gradlew", "cd ", "cp ", "mv ", "ls ", "rm ", "mkdir",
    "node ", "python3 ", "npm ", "git ", "cat ", "echo ", "EOF",
    "{", "}", '"', "[",
)


def is_code_para(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if t.startswith(CMD_STARTS):
        return True
    if t.endswith("\\"):
        return True
    if re.search(r"(\.jar|\.war|\.java|\.xml|\.html|\.toml|\.json|\.properties|\.kts)\b", t):
        if len(t) > 20 and not re.search(r"[，。：；！？]", t):
            return True
    if t.count("/") >= 2 and not re.search(r"[，。：；！？]", t) and len(t) > 15:
        return True
    if re.search(r"\s--\w+", t):
        return True
    return False


def step1_landscape_to_portrait(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    for section in doc.sections:
        new_w = section.page_height
        new_h = section.page_width
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = new_w
        section.page_height = new_h
    doc.save(str(docx_path))


def step2_fix_alignment(docx_path: Path) -> None:
    tmp = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(str(docx_path), "r") as zin:
        with zipfile.ZipFile(str(tmp), "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)
                if item == "word/styles.xml":
                    txt = data.decode("utf-8")
                    txt = txt.replace('<w:jc w:val="both"/>', '<w:jc w:val="left"/>')
                    data = txt.encode("utf-8")
                zout.writestr(item, data)
    shutil.move(str(tmp), str(docx_path))


def step3_restyle_code_paragraphs(docx_path: Path) -> int:
    doc = Document(str(docx_path))
    count = 0
    for para in doc.paragraphs:
        if not is_code_para(para.text):
            continue
        for run in para.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), "Consolas")
            rFonts.set(qn("w:hAnsi"), "Consolas")
            rFonts.set(qn("w:cs"), "Consolas")
        pPr = para._element.get_or_add_pPr()
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "left")
        shd = pPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            pPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        count += 1
    doc.save(str(docx_path))
    return count


def main() -> None:
    if len(sys.argv) < 2:
        print("用法: python3 postprocess-manual.py <docx 路径>", file=sys.stderr)
        sys.exit(64)
    docx_path = Path(sys.argv[1]).resolve()
    if not docx_path.is_file():
        print(f"文件不存在: {docx_path}", file=sys.stderr)
        sys.exit(64)

    print(f"[postprocess] 处理 {docx_path}")
    step1_landscape_to_portrait(docx_path)
    print("  [1/3] 横向 -> 纵向：完成")
    step2_fix_alignment(docx_path)
    print("  [2/3] F* 样式 jc=both -> jc=left：完成")
    n = step3_restyle_code_paragraphs(docx_path)
    print(f"  [3/3] 代码段重设字体（Consolas + 灰底）：{n} 段")


if __name__ == "__main__":
    main()
